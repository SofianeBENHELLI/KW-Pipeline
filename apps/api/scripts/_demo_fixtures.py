"""Helpers that materialise binary demo fixtures on first run.

Plain-text demo fixtures are committed under ``apps/api/fixtures/demo/``;
the PDF and DOCX cousins are *generated* the first time the seed script
runs because committing PDF/DOCX bytes is noisy in diffs and embeds
timestamps. This module owns those generators.

Both generators are deterministic in *content* — the same call always
produces a one-page PDF with the same text and a tiny DOCX with the same
paragraphs — but the byte stream is not bit-identical run-to-run because
``fpdf2`` and ``python-docx`` both stamp creation metadata. That is fine:
duplicate detection in the demo is exercised by the
``supplier_quality_policy_v1_renamed.txt`` pair, not by the binary
fixtures.

Imports of ``fpdf`` and ``docx`` are intentionally inside the functions
so this module can be imported under environments where those packages
are absent (the seed script's ``--help`` should not crash when fpdf2
isn't installed yet). Both packages ship in the API's ``[test]`` extras.
"""

from __future__ import annotations

from pathlib import Path

DEMO_DIR = Path(__file__).resolve().parent.parent / "fixtures" / "demo"
PDF_NAME = "change_request.pdf"
DOCX_NAME = "meeting_notes.docx"


def materialise_pdf(target: Path | None = None) -> Path:
    """Write a small one-page PDF to ``target`` (default: demo dir).

    Returns the path that was written. If the file already exists the
    function is a no-op — re-running the seed script must not churn
    fixture timestamps.
    """
    from fpdf import FPDF

    out = target if target is not None else DEMO_DIR / PDF_NAME
    if out.exists():
        return out
    out.parent.mkdir(parents=True, exist_ok=True)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(
        w=0,
        h=8,
        text="Engineering Change Request CR-2026-0142",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.cell(
        w=0, h=8, text="Submitted: 2026-04-22 by R. Devereaux", new_x="LMARGIN", new_y="NEXT"
    )
    pdf.cell(
        w=0,
        h=8,
        text="Affected program: Line 3 fastener bracket assembly",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.ln(2)
    pdf.set_font("Helvetica", size=10)
    paragraphs = [
        "1. Change summary",
        (
            "Replace the legacy fastener bracket part number 4471-A with the "
            "revised fastener bracket part number 4471-B across the Line 3 "
            "production cell. The revised bracket geometry corrects the bolt-pattern "
            "fit issue identified during the Q1 engineering build review."
        ),
        "2. Affected hardware",
        (
            "The bracket change affects approximately 240 fastener assemblies "
            "currently in work-in-process on Line 3. Engineering has confirmed "
            "that all downstream torque specifications remain unchanged for the "
            "revised bracket part number."
        ),
        "3. Cross-functional approval",
        (
            "Engineering, manufacturing operations, and supplier quality have "
            "signed off on the revised bracket part number. The supplier "
            "quality engineer flagged that incoming bracket lots from the "
            "supplier must continue to be sampled per the receiving inspection "
            "AQL, consistent with the Supplier Quality Policy."
        ),
        "4. Effectivity",
        (
            "Effectivity is the next production batch on Line 3 following "
            "engineering change board approval. Assemblies already in "
            "work-in-process retain the legacy bracket; mixed-bracket "
            "assemblies are not permitted in any single production batch."
        ),
    ]
    for paragraph in paragraphs:
        pdf.multi_cell(w=0, h=5, text=paragraph)
        pdf.ln(1)
    pdf.output(str(out))
    return out


def materialise_docx(target: Path | None = None) -> Path:
    """Write a small DOCX with a few paragraphs to ``target``.

    Returns the path that was written. No-op when the file already
    exists, same rationale as :func:`materialise_pdf`.
    """
    from docx import Document

    out = target if target is not None else DEMO_DIR / DOCX_NAME
    if out.exists():
        return out
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Weekly Quality Review Meeting — 2026-04-15", level=1)
    doc.add_paragraph(
        "Meeting type: weekly quality review meeting, manufacturing operations track."
    )
    doc.add_paragraph(
        "Attendees: J. Pak (chair, quality director), M. Ortega (manufacturing "
        "operations), R. Devereaux (engineering), S. Cho (supplier quality)."
    )
    doc.add_paragraph(
        "Agenda item 1 — Open non-conformance reports. The team reviewed seven "
        "open non-conformance reports from supplier lots, down from eleven the "
        "previous week. Two Major findings remain past the ten-day containment "
        "window from the Supplier Quality Policy and have been escalated to the "
        "supplier business review."
    )
    doc.add_paragraph(
        "Agenda item 2 — Receiving inspection trend. Inbound inspection AQL "
        "trend on lot 4471 series shipments continues to trend toward AQL 1.5; "
        "the supplier quality engineer recommended tightening the receiving "
        "inspection plan to AQL 1.5 ahead of the policy revision."
    )
    doc.add_paragraph(
        "Agenda item 3 — Engineering change linkage. Engineering provided an "
        "update on engineering change request CR-2026-0142, which retires the "
        "legacy fastener bracket part number 4471-A in favour of the revised "
        "fastener bracket part number 4471-B on Line 3. The change board is "
        "expected to approve the bracket change next week."
    )
    doc.add_paragraph(
        "Action items. Action 1: supplier quality to publish a revised "
        "non-conformance report containment standard. Action 2: manufacturing "
        "operations to confirm Line 3 readiness for the revised fastener "
        "bracket part number. Action 3: quality director to circulate the "
        "updated Supplier Quality Policy draft for review."
    )
    doc.save(str(out))
    return out


def materialise_all() -> list[Path]:
    """Generate every binary demo fixture that isn't already on disk."""
    return [materialise_pdf(), materialise_docx()]
