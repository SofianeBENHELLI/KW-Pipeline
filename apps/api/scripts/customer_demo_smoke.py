#!/usr/bin/env python3
"""Run the customer demo smoke path against the local Harvester API.

The runner intentionally drives the FastAPI HTTP routes via TestClient rather
than calling services directly. That keeps the demo aligned with upload
guardrails, parser dispatch, paginated catalog reads, semantic generation, and
review endpoints without requiring a live uvicorn process.
"""

from __future__ import annotations

import argparse
import io
import json
import os
import shutil
import sys
import tempfile
import zipfile
from collections.abc import Callable
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

API_ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = API_ROOT.parents[1]
if str(API_ROOT) not in sys.path:
    sys.path.insert(0, str(API_ROOT))

from app.main import create_app  # noqa: E402
from app.services.parsers.docx import DOCX_CONTENT_TYPE  # noqa: E402

TEXT_CONTENT_TYPE = "text/plain"
DEFAULT_FIXTURE_DIR = API_ROOT / "fixtures" / "customer_demo"
DEFAULT_DEMO_ROOT = REPO_ROOT / ".kw-pipeline" / "customer-demo"
DEFAULT_DATA_DIR = DEFAULT_DEMO_ROOT / "data"
DEFAULT_ARTIFACT_DIR = DEFAULT_DEMO_ROOT / "artifacts"
DEMO_DOCX_TIMESTAMP = datetime(2026, 4, 30, 12, 0, 0)
DEMO_DOCX_ZIP_TIMESTAMP = (2026, 4, 30, 12, 0, 0)


def run_customer_demo(
    *,
    fixture_dir: Path | str = DEFAULT_FIXTURE_DIR,
    data_dir: Path | str = DEFAULT_DATA_DIR,
    artifact_dir: Path | str = DEFAULT_ARTIFACT_DIR,
    reset: bool = False,
    emit: Callable[[str], None] | None = print,
) -> dict[str, Any]:
    """Run upload -> extraction -> semantic -> review for customer fixtures."""
    fixture_dir = Path(fixture_dir)
    data_dir = Path(data_dir)
    artifact_dir = Path(artifact_dir)
    _prepare_output_dirs(data_dir=data_dir, artifact_dir=artifact_dir, reset=reset)
    previous_allowlist = _ensure_demo_content_types()
    previous_kg = _ensure_knowledge_layer_enabled()
    previous_inline = _ensure_inline_extraction()

    summary: dict[str, Any] = {
        "fixture_dir": str(fixture_dir),
        "data_dir": str(data_dir),
        "artifact_dir": str(artifact_dir),
        "processed_versions": [],
    }

    try:
        return _run_demo_inside_client(
            fixture_dir=fixture_dir,
            data_dir=data_dir,
            artifact_dir=artifact_dir,
            summary=summary,
            emit=emit,
        )
    finally:
        _restore_content_types(previous_allowlist)
        _restore_env("KW_KNOWLEDGE_LAYER_ENABLED", previous_kg)
        _restore_env("KW_EXTRACTION_INLINE", previous_inline)


def _run_demo_inside_client(
    *,
    fixture_dir: Path,
    data_dir: Path,
    artifact_dir: Path,
    summary: dict[str, Any],
    emit: Callable[[str], None] | None,
) -> dict[str, Any]:
    with TestClient(create_app(persistent=True, data_dir=str(data_dir))) as client:
        _expect(client.get("/health"), "health check")
        _emit(emit, "Health check passed")

        supplier_v1 = _upload_fixture(
            client=client,
            fixture_path=fixture_dir / "acme_supplier_onboarding_policy_v1.txt",
            filename="acme_supplier_onboarding_policy_v1.txt",
            content_type=TEXT_CONTENT_TYPE,
        )
        summary["processed_versions"].append(
            _extract_generate_preview_and_validate(
                client=client,
                version=supplier_v1,
                artifact_dir=artifact_dir,
                key="supplier_policy_v1",
                reviewer_note="Validated v1 lineage and Markdown preview for demo.",
            )
        )
        _emit(emit, f"Validated supplier policy v1: {supplier_v1['sha256']}")

        # Hero document for the Demo KG (#147): repeats supplier /
        # ISO 9001 / audit / corrective-action / renewal-risk concepts
        # across enough lines to produce visible chunks, topics, and
        # chunk-to-chunk semantic edges.
        hero = _upload_fixture(
            client=client,
            fixture_path=fixture_dir / "acme_quality_program_handbook.txt",
            filename="acme_quality_program_handbook.txt",
            content_type=TEXT_CONTENT_TYPE,
        )
        summary["processed_versions"].append(
            _extract_generate_preview_and_validate(
                client=client,
                version=hero,
                artifact_dir=artifact_dir,
                key="quality_program_handbook",
                reviewer_note="Validated quality program handbook for KG demo.",
            )
        )
        _emit(emit, f"Validated quality program handbook: {hero['sha256']}")

        supplier_v2 = _upload_fixture(
            client=client,
            fixture_path=fixture_dir / "acme_supplier_onboarding_policy_v2.txt",
            filename="acme_supplier_onboarding_policy_v2.txt",
            content_type=TEXT_CONTENT_TYPE,
            document_id=supplier_v1["document_id"],
        )
        if supplier_v2["version_number"] != 2:
            raise AssertionError("Supplier policy v2 did not append as version 2.")
        summary["processed_versions"].append(
            _extract_generate_preview_and_validate(
                client=client,
                version=supplier_v2,
                artifact_dir=artifact_dir,
                key="supplier_policy_v2",
                reviewer_note="Validated v2 lineage and changed supplier requirements.",
            )
        )
        _emit(emit, f"Validated supplier policy v2: {supplier_v2['sha256']}")

        success_brief = _upload_fixture(
            client=client,
            fixture_path=fixture_dir / "customer_success_brief.txt",
            filename="customer_success_brief.txt",
            content_type=TEXT_CONTENT_TYPE,
        )
        summary["processed_versions"].append(
            _extract_generate_preview_and_validate(
                client=client,
                version=success_brief,
                artifact_dir=artifact_dir,
                key="customer_success_brief",
                reviewer_note="Validated renewal brief semantic extraction for demo.",
            )
        )
        _emit(emit, f"Validated customer success brief: {success_brief['sha256']}")

        docx_filename, docx_content = _load_docx_fixture(
            fixture_dir / "acme_contract_review_memo.json"
        )
        contract_memo = _upload_bytes(
            client=client,
            content=docx_content,
            filename=docx_filename,
            content_type=DOCX_CONTENT_TYPE,
        )
        summary["processed_versions"].append(
            _extract_generate_preview_and_validate(
                client=client,
                version=contract_memo,
                artifact_dir=artifact_dir,
                key="contract_review_memo",
                reviewer_note="Validated DOCX parser output and review transition.",
            )
        )
        _emit(emit, f"Validated contract review DOCX: {contract_memo['sha256']}")

        duplicate = _upload_fixture(
            client=client,
            fixture_path=fixture_dir / "acme_supplier_onboarding_policy_v1.txt",
            filename="archived_supplier_policy_duplicate.txt",
            content_type=TEXT_CONTENT_TYPE,
        )
        if duplicate["status"] != "DUPLICATE_DETECTED":
            raise AssertionError(f"Expected duplicate status, got {duplicate['status']}.")
        if duplicate["duplicate_of_version_id"] != supplier_v1["id"]:
            raise AssertionError("Duplicate upload did not point at supplier policy v1.")
        duplicate_extract = client.post(
            f"/documents/{duplicate['document_id']}/versions/{duplicate['id']}/extract"
        )
        if duplicate_extract.status_code != 409:
            raise AssertionError(
                "Duplicate extraction should return 409, "
                f"got {duplicate_extract.status_code}: {duplicate_extract.text}"
            )
        summary["duplicate"] = {
            "document_id": duplicate["document_id"],
            "version_id": duplicate["id"],
            "status": duplicate["status"],
            "duplicate_of_version_id": duplicate["duplicate_of_version_id"],
            "extract_status_code": duplicate_extract.status_code,
        }
        _emit(emit, "Duplicate detection path passed")

        catalog = _walk_catalog(client=client, limit=2)
        _write_json(artifact_dir / "catalog.json", {"items": catalog})
        summary["catalog_document_count"] = len(catalog)
        summary["catalog_artifact"] = str(artifact_dir / "catalog.json")

    # Aggregate graph stats across every validated version so a presenter
    # can read ``run_summary.json`` and see "13 chunks, 3 topics, 24
    # relations" without opening individual artifacts.
    summary["graph"] = _aggregate_graph_stats(summary["processed_versions"])

    _write_json(artifact_dir / "run_summary.json", summary)
    _emit(emit, f"Wrote demo artifacts to {artifact_dir}")
    return summary


def _extract_generate_preview_and_validate(
    *,
    client: TestClient,
    version: dict[str, Any],
    artifact_dir: Path,
    key: str,
    reviewer_note: str,
) -> dict[str, Any]:
    document_id = version["document_id"]
    version_id = version["id"]

    extraction = _expect(
        client.post(f"/documents/{document_id}/versions/{version_id}/extract"),
        f"extract {version['filename']}",
    )
    if not extraction["source_references"]:
        raise AssertionError(f"{version['filename']} produced no source lineage.")
    _write_json(artifact_dir / "extraction" / f"{key}.json", extraction)

    semantic = _expect(
        client.post(f"/documents/{document_id}/versions/{version_id}/semantic"),
        f"semantic {version['filename']}",
    )
    if semantic["validation_status"] != "needs_review":
        raise AssertionError("Semantic extraction did not enter needs_review.")
    if not semantic["source_references"]:
        raise AssertionError(f"{version['filename']} semantic output has no lineage.")
    _write_json(artifact_dir / "semantic" / f"{key}.needs_review.json", semantic)

    markdown_response = client.get(f"/documents/{document_id}/versions/{version_id}/markdown")
    if markdown_response.status_code != 200:
        raise AssertionError(
            f"Markdown preview failed with {markdown_response.status_code}: "
            f"{markdown_response.text}"
        )
    markdown = markdown_response.text
    if not markdown.startswith("---\n") or "## Source Lineage" not in markdown:
        raise AssertionError("Markdown preview is missing frontmatter or source lineage.")
    markdown_path = artifact_dir / "markdown" / f"{key}.md"
    markdown_path.write_text(markdown, encoding="utf-8")

    reviewed_semantic = _expect(
        client.post(
            f"/documents/{document_id}/versions/{version_id}/validate",
            json={"reviewer_note": reviewer_note},
        ),
        f"validate {version['filename']}",
    )
    if reviewed_semantic["validation_status"] != "validated":
        raise AssertionError("Review endpoint did not mark semantic output validated.")
    _write_json(artifact_dir / "semantic" / f"{key}.validated.json", reviewed_semantic)

    document = _expect(client.get(f"/documents/{document_id}"), f"document {document_id}")
    reviewed_version = next(v for v in document["versions"] if v["id"] == version_id)
    if reviewed_version["status"] != "VALIDATED":
        raise AssertionError(f"Expected VALIDATED status, got {reviewed_version['status']}.")

    # Demo KG (#145): export the graph projection alongside the
    # extraction / semantic / Markdown artifacts so reviewers can
    # eyeball chunks/topics/relations without spinning up the UI.
    graph = _expect(
        client.get(f"/documents/{document_id}/graph"),
        f"graph {version['filename']}",
    )
    _write_json(artifact_dir / "graph" / f"{key}.json", graph)
    graph_stats = _summarise_graph(graph)

    return {
        "key": key,
        "document_id": document_id,
        "version_id": version_id,
        "filename": version["filename"],
        "version_number": version["version_number"],
        "sha256": version["sha256"],
        "parser_name": extraction["parser_name"],
        "source_reference_count": len(extraction["source_references"]),
        "preview_validation_status": semantic["validation_status"],
        "review_status": reviewed_version["status"],
        "semantic_validation_status": reviewed_semantic["validation_status"],
        "markdown_artifact": str(markdown_path),
        "graph_artifact": str(artifact_dir / "graph" / f"{key}.json"),
        "graph_stats": graph_stats,
    }


# Edge kinds that count as deterministic chunk-to-chunk semantic
# relations. Mirrors lane B's relation service output. Kept module-
# local to avoid a cross-package import in the script harness.
_SEMANTIC_RELATION_KINDS: frozenset[str] = frozenset(
    {"related_to", "shares_keyword", "same_topic_as"}
)


def _aggregate_graph_stats(processed: list[dict[str, Any]]) -> dict[str, int]:
    """Sum per-version graph stats into a run-level snapshot."""
    keys = ("node_count", "edge_count", "chunk_count", "topic_count", "relation_count")
    totals = dict.fromkeys(keys, 0)
    for entry in processed:
        stats = entry.get("graph_stats") or {}
        for key in keys:
            totals[key] += int(stats.get(key, 0))
    return totals


def _summarise_graph(graph: dict[str, Any]) -> dict[str, int]:
    """Count node kinds and relation edges for ``run_summary.json``.

    Lane C's #146 smoke assertions read these counters; keep the keys
    stable so a presenter can grep for ``"chunk_count": 0`` and
    immediately spot a missing projection stage.
    """
    nodes = graph.get("nodes") or []
    edges = graph.get("edges") or []
    chunk_count = sum(1 for n in nodes if n.get("kind") == "chunk")
    topic_count = sum(1 for n in nodes if n.get("kind") == "topic")
    relation_count = sum(1 for e in edges if e.get("kind") in _SEMANTIC_RELATION_KINDS)
    return {
        "node_count": len(nodes),
        "edge_count": len(edges),
        "chunk_count": chunk_count,
        "topic_count": topic_count,
        "relation_count": relation_count,
    }


def _upload_fixture(
    *,
    client: TestClient,
    fixture_path: Path,
    filename: str,
    content_type: str,
    document_id: str | None = None,
) -> dict[str, Any]:
    return _upload_bytes(
        client=client,
        content=fixture_path.read_bytes(),
        filename=filename,
        content_type=content_type,
        document_id=document_id,
    )


def _upload_bytes(
    *,
    client: TestClient,
    content: bytes,
    filename: str,
    content_type: str,
    document_id: str | None = None,
) -> dict[str, Any]:
    params = {"document_id": document_id} if document_id is not None else None
    version = _expect(
        client.post(
            "/documents/upload",
            params=params,
            files={"file": (filename, content, content_type)},
        ),
        f"upload {filename}",
    )
    if len(version["sha256"]) != 64:
        raise AssertionError(f"{filename} did not return a SHA-256 digest.")
    return version


def _walk_catalog(*, client: TestClient, limit: int) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    cursor: str | None = None
    while True:
        params: dict[str, Any] = {"limit": limit}
        if cursor is not None:
            params["cursor"] = cursor
        page = _expect(client.get("/documents", params=params), "catalog page")
        items.extend(page["items"])
        cursor = page["next_cursor"]
        if cursor is None:
            return items


def _load_docx_fixture(spec_path: Path) -> tuple[str, bytes]:
    spec = json.loads(spec_path.read_text(encoding="utf-8"))
    document = DocxDocument()
    properties = document.core_properties
    properties.author = "KW Pipeline"
    properties.created = DEMO_DOCX_TIMESTAMP
    properties.last_modified_by = "KW Pipeline"
    properties.modified = DEMO_DOCX_TIMESTAMP
    properties.revision = 1
    properties.title = spec["filename"]
    for paragraph in spec["paragraphs"]:
        document.add_paragraph(paragraph)
    for rows in spec.get("tables", []):
        if not rows:
            continue
        table = document.add_table(rows=len(rows), cols=len(rows[0]))
        for row_index, row in enumerate(rows):
            for col_index, cell_text in enumerate(row):
                table.rows[row_index].cells[col_index].text = cell_text
    buffer = io.BytesIO()
    document.save(buffer)
    return spec["filename"], _canonicalize_docx(buffer.getvalue())


def _canonicalize_docx(content: bytes) -> bytes:
    source = io.BytesIO(content)
    target = io.BytesIO()
    with (
        zipfile.ZipFile(source, "r") as input_zip,
        zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as output_zip,
    ):
        for name in sorted(input_zip.namelist()):
            input_info = input_zip.getinfo(name)
            output_info = zipfile.ZipInfo(
                filename=name,
                date_time=DEMO_DOCX_ZIP_TIMESTAMP,
            )
            output_info.compress_type = zipfile.ZIP_DEFLATED
            output_info.external_attr = input_info.external_attr
            output_zip.writestr(output_info, input_zip.read(name))
    return target.getvalue()


def _ensure_demo_content_types() -> str | None:
    """Widen the upload allowlist so the smoke run accepts text+DOCX, returning
    the previous value so the caller can restore it. Returning ``None`` means
    the env var was unset before this call.

    The smoke runner is sometimes invoked from a pytest process; permanently
    mutating ``os.environ`` would leak into sibling tests that assert against
    the default allowlist (e.g. ``test_disallowed_content_type_returns_415``).
    The caller pairs this with a ``finally`` to restore.
    """
    previous = os.environ.get("ALLOWED_CONTENT_TYPES")
    raw = previous or ""
    allowed = {entry.strip() for entry in raw.split(",") if entry.strip()}
    allowed.update({TEXT_CONTENT_TYPE, DOCX_CONTENT_TYPE})
    os.environ["ALLOWED_CONTENT_TYPES"] = ",".join(sorted(allowed))
    return previous


def _restore_content_types(previous: str | None) -> None:
    if previous is None:
        os.environ.pop("ALLOWED_CONTENT_TYPES", None)
    else:
        os.environ["ALLOWED_CONTENT_TYPES"] = previous


def _ensure_inline_extraction() -> str | None:
    """Pin the smoke run to inline extraction (ADR-006 / PR-3 fallback).

    The smoke runner asserts ``POST /documents/.../extract`` returns
    HTTP 200 with a :class:`RawExtraction` body. PR-3 flipped the
    production default to async (202 + ``ExtractionJobSnapshot``), so
    we explicitly re-enable the legacy synchronous path here. Same
    caller-pair pattern as :func:`_ensure_knowledge_layer_enabled` so
    we don't leak into sibling tests that expect the new async
    default.
    """
    previous = os.environ.get("KW_EXTRACTION_INLINE")
    os.environ["KW_EXTRACTION_INLINE"] = "true"
    return previous


def _ensure_knowledge_layer_enabled() -> str | None:
    """Turn on the v0.2 KG projection for the smoke run.

    The runner ships a knowledge-graph demo path (#145), so chunks /
    topics / relations must be projected even when the host process
    runs with the layer disabled by default. Same caller-pair pattern
    as :func:`_ensure_demo_content_types` so we don't leak into
    sibling tests.
    """
    previous = os.environ.get("KW_KNOWLEDGE_LAYER_ENABLED")
    os.environ["KW_KNOWLEDGE_LAYER_ENABLED"] = "true"
    return previous


def _restore_env(name: str, previous: str | None) -> None:
    if previous is None:
        os.environ.pop(name, None)
    else:
        os.environ[name] = previous


def _prepare_output_dirs(*, data_dir: Path, artifact_dir: Path, reset: bool) -> None:
    if reset:
        _safe_rmtree(data_dir)
        _safe_rmtree(artifact_dir)
    elif _has_entries(data_dir) or _has_entries(artifact_dir):
        raise RuntimeError(
            "Demo data or artifacts already exist. Re-run with --reset or choose empty paths."
        )

    for child in [
        data_dir,
        artifact_dir / "extraction",
        artifact_dir / "markdown",
        artifact_dir / "semantic",
    ]:
        child.mkdir(parents=True, exist_ok=True)


def _safe_rmtree(path: Path) -> None:
    if not path.exists():
        return
    resolved = path.resolve()
    generated_root = (REPO_ROOT / ".kw-pipeline").resolve()
    temp_root = Path(tempfile.gettempdir()).resolve()
    if resolved == generated_root:
        raise RuntimeError(f"Refusing to delete generated root directly: {resolved}")
    if not (
        generated_root in resolved.parents or resolved == temp_root or temp_root in resolved.parents
    ):
        raise RuntimeError(
            f"Refusing to reset {resolved}; use a path under {generated_root} or {temp_root}."
        )
    shutil.rmtree(resolved)


def _has_entries(path: Path) -> bool:
    return path.exists() and any(path.iterdir())


def _expect(response, label: str, expected_status: int = 200) -> Any:
    if response.status_code != expected_status:
        raise AssertionError(f"{label} failed with HTTP {response.status_code}: {response.text}")
    return response.json()


def _write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def _emit(emit: Callable[[str], None] | None, message: str) -> None:
    if emit is not None:
        emit(message)


def _path_arg(raw: str) -> Path:
    return Path(raw).expanduser()


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Run the local customer demo smoke path through Harvester API routes."
    )
    parser.add_argument(
        "--fixture-dir",
        type=_path_arg,
        default=DEFAULT_FIXTURE_DIR,
        help=f"Fixture directory to upload. Default: {DEFAULT_FIXTURE_DIR}",
    )
    parser.add_argument(
        "--data-dir",
        type=_path_arg,
        default=DEFAULT_DATA_DIR,
        help=f"Persistent API data directory. Default: {DEFAULT_DATA_DIR}",
    )
    parser.add_argument(
        "--artifact-dir",
        type=_path_arg,
        default=DEFAULT_ARTIFACT_DIR,
        help="Output directory for catalog, extraction, semantic JSON, and Markdown.",
    )
    parser.add_argument(
        "--reset",
        action="store_true",
        help="Delete the selected demo data and artifact directories before running.",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    summary = run_customer_demo(
        fixture_dir=args.fixture_dir,
        data_dir=args.data_dir,
        artifact_dir=args.artifact_dir,
        reset=args.reset,
    )
    print("")
    print("Customer demo smoke complete")
    print(f"Validated versions: {len(summary['processed_versions'])}")
    print(f"Catalog documents: {summary['catalog_document_count']}")
    print(f"Artifacts: {summary['artifact_dir']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
